"""
Word Document Acronym Finder

This script analyzes Word documents to identify potential acronyms,
their definitions, and locations. It can optionally use a predefined CSV file
of known acronyms and their definitions.

Requirements:
    python-docx>=0.8.11
    pandas>=1.3.0
"""

import logging
import re
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Set, Tuple

import pandas as pd
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import _Cell, Table
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class AcronymFinder:
    """Main class for finding and processing acronyms in Word documents."""
    
    def __init__(self, docx_path: str, known_acronyms_csv: Optional[str] = None,
                 log_level: int = logging.INFO) -> None:
        """
        Initialize the AcronymFinder.

        Args:
            docx_path: Path to the Word document
            known_acronyms_csv: Optional path to CSV file with known acronyms
            log_level: Logging level (default: INFO)
        """
        self.docx_path = Path(docx_path)
        self.setup_logging(log_level)
        self.known_acronyms = self._load_known_acronyms(known_acronyms_csv) if known_acronyms_csv else {}
        self.found_acronyms: Dict[str, Dict] = {}
        
    def setup_logging(self, log_level: int) -> None:
        """Set up logging with timestamp-based filename."""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        log_file = f"acronym_finder_{timestamp}.log"
        
        # Create logger
        self.logger = logging.getLogger(__name__)
        self.logger.setLevel(log_level)
        
        # Remove any existing handlers
        self.logger.handlers = []
        
        # Create handlers
        file_handler = logging.FileHandler(log_file)
        console_handler = logging.StreamHandler()
        
        # Create formatter
        formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
        
        # Set formatter for handlers
        file_handler.setFormatter(formatter)
        console_handler.setFormatter(formatter)
        
        # Add handlers to logger
        self.logger.addHandler(file_handler)
        self.logger.addHandler(console_handler)
        
    def _load_known_acronyms(self, csv_path: str) -> Dict[str, str]:
        """
        Load known acronyms from CSV file.
        
        Args:
            csv_path: Path to CSV file containing acronyms and definitions
            
        Returns:
            Dictionary mapping acronyms to their definitions
        """
        try:
            df = pd.read_csv(csv_path)
            return dict(zip(df['Acronym'], df['Definition']))
        except Exception as e:
            self.logger.error(f"Error loading known acronyms: {e}")
            return {}
            
    def _is_potential_acronym(self, word: str) -> bool:
        """
        Check if a word could be an acronym based on patterns.
        
        Args:
            word: Word to check
            
        Returns:
            Boolean indicating if the word might be an acronym
        """
        # Same acronym patterns as before
        acronym_patterns = [
            r'^[A-Z0-9][A-Z0-9]{1,5}$',
            r'^[A-Z][&][A-Z]$',
            r'^[A-Z]+/[A-Z]+$',
            r'^[A-Z0-9]+-[A-Z0-9]+$'
        ]
        
        exclusions = {'I', 'A', 'OK', 'ID', 'NO', 'AM', 'PM', 'THE'}
        
        if word in exclusions:
            return False
            
        return any(bool(re.match(pattern, word)) for pattern in acronym_patterns)
    
    def _find_potential_definition(self, text: str, acronym: str) -> Optional[str]:
        """
        Look for potential definition of an acronym in surrounding text.
        
        Args:
            text: Text to search for definition
            acronym: Acronym to find definition for
            
        Returns:
            Potential definition if found, None otherwise
        """
        escaped_acronym = re.escape(acronym)
        
        patterns = [
            rf'{escaped_acronym}\s*\(([\w\s,/-]+)\)',
            rf'\(([\w\s,/-]+)\)\s*{escaped_acronym}',
            rf'{escaped_acronym}:\s*([\w\s,/-]+)',
            rf'{escaped_acronym}\s*-\s*([\w\s,/-]+)',
            rf'{escaped_acronym}\s+stands\s+for\s+([\w\s,/-]+)',
            rf'{escaped_acronym}\s+means\s+([\w\s,/-]+)'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        return None

    def _process_table_cell(self, cell: _Cell) -> str:
        """
        Extract text from a table cell.
        
        Args:
            cell: Word table cell object
            
        Returns:
            Text content of the cell
        """
        try:
            return cell.text.strip()
        except Exception as e:
            self.logger.warning(f"Error extracting text from table cell: {e}")
            return ""

    def _process_table(self, table: Table) -> str:
        """
        Process a Word table and extract all text.
        
        Args:
            table: Word table object
            
        Returns:
            Concatenated text from all cells
        """
        table_text = []
        
        try:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = self._process_table_cell(cell)
                    if cell_text:
                        table_text.append(cell_text)
            
            return " ".join(table_text)
        except Exception as e:
            self.logger.warning(f"Error processing table: {e}")
            return ""
    
    def process_paragraph(self, paragraph: Paragraph, page_number: int) -> None:
        """
        Process a single paragraph to find acronyms and their definitions.
        
        Args:
            paragraph: Word paragraph object
            page_number: Number of the current page
        """
        text = paragraph.text
        words = re.findall(r'\b[\w/&-]+\b', text)
        
        for word in words:
            if self._is_potential_acronym(word):
                if word not in self.found_acronyms:
                    self.found_acronyms[word] = {
                        'definition': self.known_acronyms.get(word),
                        'pages': set()
                    }
                
                self.found_acronyms[word]['pages'].add(page_number)
                
                # Look for definition if not already known
                if not self.found_acronyms[word]['definition']:
                    definition = self._find_potential_definition(text, word)
                    if definition:
                        self.found_acronyms[word]['definition'] = definition
                        self.logger.info(f"Found definition for {word}: {definition}")
    
    def create_acronym_table_document(self) -> None:
        """
        Create a new document with the acronym table.
        """
        doc = Document()
        
        # Add title
        title = doc.add_paragraph("Acronyms Found")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.style = 'Heading 1'
        
        # Create table
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        table.autofit = True
        
        # Set column widths
        for column in table.columns:
            for cell in column.cells:
                cell.width = Inches(2)
        
        # Set headers
        headers = table.rows[0].cells
        headers[0].text = 'Acronym'
        headers[1].text = 'Definition'
        headers[2].text = 'Page Numbers'
        
        # Populate table
        for acronym, info in sorted(self.found_acronyms.items()):
            row = table.add_row().cells
            row[0].text = acronym
            row[1].text = info['definition'] or 'Unknown'
            row[2].text = ', '.join(map(str, sorted(info['pages'])))
        
        # Save the document
        output_path = self.docx_path.with_name(f"{self.docx_path.stem}_acronyms.docx")
        doc.save(output_path)
        self.logger.info(f"Saved acronym table to {output_path}")
    
    def process_document(self) -> None:
        """Process the entire document to find and document acronyms."""
        try:
            doc = Document(self.docx_path)
            current_page = 1
            chars_on_page = 0
            chars_per_page = 1800  # Approximate characters per page
            
            # Process paragraphs
            for paragraph in doc.paragraphs:
                chars_on_page += len(paragraph.text)
                if chars_on_page > chars_per_page:
                    current_page += 1
                    chars_on_page = len(paragraph.text)
                self.process_paragraph(paragraph, current_page)
            
            # Process tables
            for table in doc.tables:
                table_text = self._process_table(table)
                chars_on_page += len(table_text)
                if chars_on_page > chars_per_page:
                    current_page += 1
                    chars_on_page = len(table_text)
                
                # Process table text directly without creating temporary paragraph
                text = table_text
                words = re.findall(r'\b[\w/&-]+\b', text)
                
                for word in words:
                    if self._is_potential_acronym(word):
                        if word not in self.found_acronyms:
                            self.found_acronyms[word] = {
                                'definition': self.known_acronyms.get(word),
                                'pages': set()
                            }
                        
                        self.found_acronyms[word]['pages'].add(current_page)
                        
                        # Look for definition if not already known
                        if not self.found_acronyms[word]['definition']:
                            definition = self._find_potential_definition(text, word)
                            if definition:
                                self.found_acronyms[word]['definition'] = definition
                                self.logger.info(f"Found definition for {word}: {definition}")
            
            # Create acronym table in a separate document
            self.create_acronym_table_document()
            
        except Exception as e:
            self.logger.error(f"Error processing document: {e}")
            raise

def main():
    """Main entry point for the script."""
    import argparse
    
    parser = argparse.ArgumentParser(description="Find acronyms in Word documents")
    parser.add_argument("docx_path", help="Path to the Word document")
    parser.add_argument("--known-acronyms", help="Path to CSV file with known acronyms")
    parser.add_argument("--log-level", default="INFO",
                      choices=["DEBUG", "INFO", "WARNING", "ERROR", "CRITICAL"],
                      help="Set the logging level")
    
    args = parser.parse_args()
    
    finder = AcronymFinder(
        args.docx_path,
        args.known_acronyms,
        getattr(logging, args.log_level)
    )
    finder.process_document()

if __name__ == "__main__":
    main()
