import re
import pandas as pd
from typing import Dict, Optional, Set, List
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import _Cell, Table
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from utils.logger import setup_logger
from utils.progress import ProgressTracker

logger = setup_logger(__name__)

class AcronymAnalyzer:
    """Analyzes documents for acronyms and their definitions."""
    
    def __init__(self, docx_path: str, known_acronyms_csv: Optional[str] = None,
                 excluded_acronyms_csv: Optional[str] = None):
        """Initialize the acronym analyzer."""
        self.docx_path = Path(docx_path)
        if not self.docx_path.exists():
            raise FileNotFoundError(f"Document not found: {docx_path}")
            
        # Initialize document tracking attributes
        self.current_page = 1
        self.chars_on_page = 0
        self.chars_per_page = 1800  # Approximate characters per page
        
        # Load acronyms and exclusions
        self.known_acronyms = self._load_known_acronyms(known_acronyms_csv) if known_acronyms_csv else {}
        self.excluded_acronyms = self._load_excluded_acronyms(excluded_acronyms_csv) if excluded_acronyms_csv else set()
        self.found_acronyms: Dict[str, Dict] = {}
        
        # Initialize found_acronyms with known acronyms
        for acronym, definition in self.known_acronyms.items():
            if acronym not in self.excluded_acronyms:
                acronym = str(acronym).strip()
                if acronym and acronym.lower() != 'nan':  # Additional check to exclude 'nan' strings
                    self.found_acronyms[acronym] = {
                        'definition': definition.strip() if definition else None,
                        'pages': set()
                    }
        
        # Set up acronym patterns with at least one uppercase letter
        self.acronym_patterns = [
            r'^(?=.*[A-Z])[A-Z0-9]{2,6}$',  # Basic acronyms (2-6 chars, at least one letter)
            r'^[A-Z][&][A-Z]$',             # Special case for X&Y format
            r'^(?=.*[A-Z])[A-Z]+/[A-Z]+$',  # Slash-separated acronyms (at least one letter)
            r'^(?=.*[A-Z])[A-Z0-9]+-[A-Z0-9]+$'  # Hyphenated acronyms (at least one letter)
        ]
        
        # Set up default exclusions
        self.default_exclusions = {'I', 'A', 'OK', 'ID', 'NO', 'AM', 'PM', 'THE'}
        
        logger.info(f"Acronym analyzer initialized for {docx_path}")
        logger.debug(f"Known acronyms: {len(self.known_acronyms)}")
        logger.debug(f"Excluded acronyms: {len(self.excluded_acronyms)}")

    def _load_known_acronyms(self, csv_path: str) -> Dict[str, str]:
        """
        Load known acronyms from CSV file.
    
        Args:
            csv_path: Path to CSV file containing acronyms and definitions
    
        Returns:
            Dictionary mapping acronyms to their definitions
        """
        try:
            df = pd.read_csv(csv_path, dtype={'Acronym': str, 'Definition': str})
            # Strip whitespace
            df['Acronym'] = df['Acronym'].str.strip()
            df['Definition'] = df['Definition'].str.strip()
            # Drop rows with missing or empty 'Acronym'
            initial_count = len(df)
            df = df[df['Acronym'].notna() & (df['Acronym'] != '')]
            dropped_count = initial_count - len(df)
            if dropped_count > 0:
                logger.warning(f"Dropped {dropped_count} rows with missing or empty 'Acronym' in {csv_path}")
            
            if 'Acronym' in df.columns and 'Definition' in df.columns:
                acronyms = dict(zip(df['Acronym'], df['Definition']))
                logger.info(f"Loaded {len(acronyms)} known acronyms from {csv_path}")
                return acronyms
            else:
                logger.error(f"CSV file {csv_path} missing required columns (Acronym, Definition)")
                return {}
        except Exception as e:
            logger.error(f"Error loading known acronyms: {e}")
            return {}

    def _load_excluded_acronyms(self, csv_path: str) -> Set[str]:
        """
        Load acronyms to exclude from CSV file.
    
        Args:
            csv_path: Path to CSV file containing acronyms to exclude
    
        Returns:
            Set of acronyms to exclude from analysis
        """
        try:
            df = pd.read_csv(csv_path, dtype={'Exclusion': str})
            # Strip whitespace and drop NaN
            df['Exclusion'] = df['Exclusion'].str.strip()
            exclusions = set(df['Exclusion'].dropna())
            logger.info(f"Loaded {len(exclusions)} excluded acronyms from {csv_path}")
            return exclusions
        except Exception as e:
            logger.error(f"Error loading excluded acronyms: {e}")
            return set()

    def _is_potential_acronym(self, word: str) -> bool:
        """
        Check if a word could be an acronym based on patterns.
        
        Args:
            word: Word to check
            
        Returns:
            Boolean indicating if the word might be an acronym
        """
        # Check exclusions first
        if word in self.default_exclusions or word in self.excluded_acronyms:
            return False
        
        # Exclude purely numeric words
        if re.fullmatch(r'\d+', word):
            logger.debug(f"Excluded purely numeric word: {word}")
            return False
        
        # Exclude numeric words with dashes or slashes (e.g., "0-1", "1/2")
        if re.fullmatch(r'(\d+[-/])+(\d+)', word):
            logger.debug(f"Excluded numeric word with separators: {word}")
            return False
        
        # Check against acronym patterns (ensures at least one letter due to patterns)
        is_match = any(bool(re.match(pattern, word)) for pattern in self.acronym_patterns)
        if is_match:
            return True
        else:
            return False

    def _find_potential_definition(self, text: str, acronym: str) -> Optional[str]:
        """
        Look for potential definition of an acronym in surrounding text.
        
        Args:
            text: Text to search for definition
            acronym: Acronym to find definition for
            
        Returns:
            Potential definition if found, None otherwise
        """
        escaped_acronym = re.escape(acronym)
        
        patterns = [
            rf'{escaped_acronym}\s*\(([\w\s,/-]+)\)',              # Acronym (Definition)
            rf'\(([\w\s,/-]+)\)\s*{escaped_acronym}',              # (Definition) Acronym
            rf'{escaped_acronym}:\s*([\w\s,/-]+)',                 # Acronym: Definition
            rf'{escaped_acronym}\s*-\s*([\w\s,/-]+)',             # Acronym - Definition
            rf'{escaped_acronym}\s+stands\s+for\s+([\w\s,/-]+)',   # Acronym stands for Definition
            rf'{escaped_acronym}\s+means\s+([\w\s,/-]+)'          # Acronym means Definition
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        return None

    def reset_page_tracking(self):
        """Reset page tracking counters."""
        self.current_page = 1
        self.chars_on_page = 0

    def update_page_tracking(self, text_length: int) -> None:
        """
        Update page tracking based on text length.
        
        Args:
            text_length: Length of text being processed
        """
        self.chars_on_page += int(text_length)  # Ensure integer
        if self.chars_on_page > self.chars_per_page:
            self.current_page += 1
            self.chars_on_page = int(text_length)  # Ensure integer

    def process_text(self, text: str, page_number: int) -> None:
        """
        Process text to find acronyms and their definitions.
        
        Args:
            text: Text to analyze
            page_number: Current page number
        """
        try:
            words = re.findall(r'\b[\w/&-]+\b', text)
            page_number = int(page_number)  # Ensure page_number is an integer
            
            for word in words:
                original_word = word  # Keep the original word for logging
                word = str(word).strip()  # Ensure word is a string and remove whitespace
                if self._is_potential_acronym(word):
                    if word and word.lower() != 'nan':  # Additional check
                        if word not in self.found_acronyms:
                            definition = self._find_potential_definition(text, word)
                            if not definition:
                                definition = self.known_acronyms.get(word)
                            
                            self.found_acronyms[word] = {
                                'definition': definition.strip() if definition else None,
                                'pages': {page_number}
                            }
                        else:
                            # Add the page number to existing pages set
                            self.found_acronyms[word]['pages'].add(page_number)
                else:
                    logger.debug(f"Word '{original_word}' is not a valid acronym and was excluded.")
            
                # Debugging: Log the type of the word
                if not isinstance(word, str):
                    logger.warning(f"Non-string acronym detected: {word} (type: {type(word)})")
        except Exception as e:
            logger.error(f"Error processing text: {e}")
            logger.debug(f"Text sample: {text[:100]}...")
            logger.debug(f"Page number type: {type(page_number)}, value: {page_number}")
            raise

    def process_paragraph(self, paragraph: Paragraph, page_number: int) -> None:
        """Process a paragraph to find acronyms."""
        try:
            self.process_text(paragraph.text, page_number)
        except Exception as e:
            logger.error(f"Error processing paragraph: {e}")
            raise

    def process_table(self, table: Table, page_number: int) -> None:
        """Process a table to find acronyms."""
        try:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self.process_text(paragraph.text, page_number)
        except Exception as e:
            logger.error(f"Error processing table: {e}")
            raise

    def create_acronym_report(self) -> Document:
        """
        Create a document containing the acronym reference table.
        
        Returns:
            Word document containing formatted acronym report
        """
        # Debugging: Check types of all acronyms
        non_string_keys = [k for k in self.found_acronyms.keys() if not isinstance(k, str)]
        if non_string_keys:
            logger.error(f"Non-string acronym keys detected: {non_string_keys}")
            # Optionally, skip these entries instead of raising an error
            # Or clean them up
            self.found_acronyms = {k: v for k, v in self.found_acronyms.items() if isinstance(k, str)}
            logger.info(f"Skipped {len(non_string_keys)} non-string acronym entries.")

        doc = Document()
        
        # Add title
        title = doc.add_paragraph("Acronyms Found")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.style = 'Heading 1'
        
        # Create table with only acronym and definition columns
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.autofit = True
        
        # Set column widths
        for column in table.columns:
            for cell in column.cells:
                cell.width = Inches(3)
        
        # Set headers
        headers = table.rows[0].cells
        headers[0].text = 'Acronym'
        headers[1].text = 'Definition'
        
        # Populate table, sorted alphabetically by acronym
        for acronym, info in sorted(self.found_acronyms.items(), key=lambda item: item[0].upper()):
            row = table.add_row().cells
            row[0].text = acronym
            row[1].text = info['definition'] or 'Unknown'
        
        return doc

    def analyze_document(self) -> Dict[str, Dict]:
        """Analyze the entire document for acronyms."""
        try:
            doc = Document(self.docx_path)
            self.reset_page_tracking()
            
            # Count total elements for progress tracking
            total_elements = len(doc.paragraphs) + len(doc.tables)
            progress = ProgressTracker(total_elements, "Analyzing acronyms")
            
            # Process paragraphs
            for paragraph in doc.paragraphs:
                text_length = len(paragraph.text)
                self.update_page_tracking(text_length)
                self.process_paragraph(paragraph, self.current_page)
                progress.update()
            
            # Process tables
            for table in doc.tables:
                table_text = "".join(
                    cell.text for row in table.rows for cell in row.cells
                )
                text_length = len(table_text)
                self.update_page_tracking(text_length)
                self.process_table(table, self.current_page)
                progress.update()
            
            progress.complete()
            
            # Create and save acronym report
            report = self.create_acronym_report()
            output_path = self.docx_path.with_name(f"{self.docx_path.stem}_acronyms.docx")
            report.save(output_path)
            logger.info(f"Saved acronym report to {output_path}")
            
            return self.found_acronyms
            
        except Exception as e:
            logger.error(f"Error analyzing document: {e}")
            raise

    def get_statistics(self) -> Dict[str, int]:
        """
        Get statistics about found acronyms.
        
        Returns:
            Dictionary containing acronym statistics
        """
        total_acronyms = len(self.found_acronyms)
        defined_acronyms = sum(
            1 for info in self.found_acronyms.values() 
            if info['definition']
        )
        undefined_acronyms = total_acronyms - defined_acronyms
        
        return {
            'total_acronyms': total_acronyms,
            'defined_acronyms': defined_acronyms,
            'undefined_acronyms': undefined_acronyms
        }
