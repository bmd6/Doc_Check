import re
from typing import List, Optional, Dict, Tuple
from docx.document import Document
from docx.text.paragraph import Paragraph
from docx.table import _Cell, Table
from docx.oxml.shared import qn
from docx.shared import Pt

from models.issues import StyleIssue
from models.fonts import FontInfo, FontUsageSummary
from utils.logger import setup_logger
from utils.progress import ProgressTracker

logger = setup_logger(__name__)

class StyleAnalyzer:
    """
    Analyzes document style elements including fonts, headers, tables, and captions.
    
    This analyzer focuses on:
    - Font consistency
    - Header formatting
    - Table styles and headers
    - Caption placement and formatting
    """
    
    def __init__(self):
        """Initialize the style analyzer."""
        self.font_usage = FontUsageSummary()
        self.current_page = 1
        logger.info("Style analyzer initialized successfully")

    def _get_font_info(self, run) -> Optional[FontInfo]:
        """
        Extract font information from a text run.
        
        Args:
            run: Document run object containing text formatting
            
        Returns:
            FontInfo object containing formatting details
        """
        try:
            font = run._element.rPr.rFonts
            size_element = run._element.rPr.sz
            bold = run._element.rPr.b
            italic = run._element.rPr.i

            # Get font name
            font_name = None
            if font is not None:
                # Try different font attributes in order of preference
                font_name = (font.get(qn('w:ascii')) or
                           font.get(qn('w:hAnsi')) or
                           font.get(qn('w:cs')) or
                           font.get(qn('w:eastAsia')))

            # Get font size
            font_size = None
            if size_element is not None:
                try:
                    font_size = float(size_element.val)
                except (ValueError, TypeError):
                    font_size = None

            return FontInfo(
                name=font_name,
                size=font_size,
                bold=bold is not None,
                italic=italic is not None
            )
        except AttributeError as e:
            logger.debug(f"AttributeError in _get_font_info: {e}")
            return FontInfo()

    def _get_paragraph_style_info(self, paragraph: Paragraph) -> str:
        """
        Get a string representation of paragraph style information.
        
        Args:
            paragraph: Paragraph to analyze
            
        Returns:
            String describing the paragraph's style
        """
        total_runs = 0
        font_counts = {}  # Track frequency of each font configuration

        for run in paragraph.runs:
            if run.text.strip():  # Only consider runs with actual text
                total_runs += 1
                current_font = self._get_font_info(run)
                font_str = str(current_font)

                if font_str not in font_counts:
                    font_counts[font_str] = 0
                font_counts[font_str] += 1

        if not font_counts:  # No text runs found
            return "Default"

        # Find the most common font configuration
        most_common_font = max(font_counts.items(), key=lambda x: x[1])
        percentage = (most_common_font[1] / total_runs) * 100

        if percentage == 100:  # All runs use the same font
            return most_common_font[0]
        else:  # Mixed fonts
            mixed_fonts = ', '.join(f"'{font}'" for font in font_counts.keys())
            return f"Mixed ({mixed_fonts})"

    def _check_table_header_repeat(self, table: Table) -> bool:
        """
        Check if table headers are set to repeat on each page.
        
        Args:
            table: Table to check
            
        Returns:
            Boolean indicating if headers repeat
        """
        try:
            for row in table.rows:
                tr = row._tr  # Access the underlying CT_Row object
                trPr = tr.trPr  # Access the row properties
                if trPr is not None:
                    tbl_header = trPr.find(qn('./w:tblHeader')) # Find tblHeader element using xpath
                    if tbl_header and tbl_header[0].get(qn('w:val')) in ('1', 'true', True):
                        return True
            return False
        except Exception as e:
            logger.debug(f"Error checking table header repeat: {e}")
            return False

    def _process_table_cell(self, cell: _Cell) -> Tuple[str, Dict[str, int]]:
        """
        Process a table cell and extract text and font information.
        
        Args:
            cell: Table cell to process
            
        Returns:
            Tuple of (cell text, font usage dictionary)
        """
        text = []
        font_usage = {}
        
        try:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.text.strip():
                        font_info = self._get_font_info(run)
                        font_str = str(font_info)
                        font_usage[font_str] = font_usage.get(font_str, 0) + len(run.text)
                        text.append(run.text)
        except Exception as e:
            logger.debug(f"Error processing table cell: {e}")
        
        return ' '.join(text), font_usage

    def analyze_paragraph(self, paragraph: Paragraph, 
                        page: int, section: Optional[str]) -> List[StyleIssue]:
        """
        Analyze a paragraph for style issues.
        
        Args:
            paragraph: Paragraph to analyze
            page: Current page number
            section: Current document section
            
        Returns:
            List of style issues found
        """
        issues = []
        style_info = self._get_paragraph_style_info(paragraph)
        char_count = len(paragraph.text)

        # Track font usage based on paragraph type
        if paragraph.style.name.startswith('Heading'):
            try:
                header_level = int(paragraph.style.name.split()[-1])
            except (ValueError, IndexError):
                header_level = 1
            self.font_usage.add_header_font(header_level, style_info, char_count)
            
            # Check header formatting
            header_font = self._get_font_info(paragraph.runs[0]) if paragraph.runs else None
            if header_font:
                expected_size = 14 - (header_level - 1)  # Example size calculation
                if header_font.size and abs(header_font.size - expected_size) > 0.5:
                    issues.append(StyleIssue(
                        type="Header Format",
                        element=f"Heading {header_level}",
                        expected=f"Font size {expected_size}pt",
                        found=f"Font size {header_font.size}pt",
                        page=page,
                        section=section,
                        context=paragraph.text[:50]
                    ))
        
        elif "Caption" in paragraph.style.name.lower():
            self.font_usage.add_caption_font(style_info, char_count)
            
            # Check caption formatting
            caption_font = self._get_font_info(paragraph.runs[0]) if paragraph.runs else None
            if caption_font and caption_font.size and caption_font.size > 11:
                issues.append(StyleIssue(
                    type="Caption Format",
                    element="Caption",
                    expected="Font size ≤ 11pt",
                    found=f"Font size {caption_font.size}pt",
                    page=page,
                    section=section,
                    context=paragraph.text[:50]
                ))
        
        elif not paragraph.style.name.startswith('TOC'):
            self.font_usage.add_body_font(style_info, char_count)
            
            # Check body text formatting
            if 'Mixed' in style_info:
                issues.append(StyleIssue(
                    type="Inconsistent Font",
                    element="Body Text",
                    expected="Consistent font formatting",
                    found=style_info,
                    page=page,
                    section=section,
                    context=paragraph.text[:50]
                ))

        return issues

    def analyze_table(self, table: Table, page: int, 
                     section: Optional[str]) -> List[StyleIssue]:
        """
        Analyze a table for style issues.
        
        Args:
            table: Table to analyze
            page: Current page number
            section: Current document section
            
        Returns:
            List of style issues found
        """
        issues = []

        # Check header rows repeat
        has_repeating_header = False
        try:
            has_repeating_header = self._check_table_header_repeat(table)
        except Exception as e:
            logger.debug(f"Error checking table headers: {e}")

        
        if not has_repeating_header and len(table.rows) > 1:
            issues.append(StyleIssue(
                type="Table Header",
                element="Table",
                expected="Headers should repeat on each page",
                found="Headers do not repeat",
                page=page,
                section=section,
                context=table.rows[0].cells[0].text[:50] if table.rows else ""
            ))

        # Process cells
        try:
            for row in table.rows:
                for cell in row.cells:
                    text, font_usage = self._process_table_cell(cell)
                    
                    # Track font usage
                    for font_str, count in font_usage.items():
                        self.font_usage.add_table_font(font_str, count)
                    
                    # Check for inconsistent formatting within cells
                    if len(font_usage) > 1:
                        issues.append(StyleIssue(
                            type="Table Cell Format",
                            element="Table Cell",
                            expected="Consistent font formatting",
                            found=f"Mixed fonts: {', '.join(font_usage.keys())}",
                            page=page,
                            section=section,
                            context=text[:50]
                        ))
        except Exception as e:
            logger.debug(f"Error processing table cells: {e}")

        return issues

    def analyze_document(self, document: Document) -> List[StyleIssue]:
        """
        Analyze entire document for style issues.
        
        Args:
            document: Document to analyze
            
        Returns:
            List of all style issues found
        """
        issues = []
        current_section = None
        
        # Track progress
        total_paragraphs = len(document.paragraphs) + sum(1 for table in document.tables)
        progress = ProgressTracker(total_paragraphs, "Analyzing document style")

        # Analyze paragraphs and tables
        for element in document.element.body:
            if isinstance(element, Paragraph):
                paragraph = Paragraph(element, document)
                
                # Update section tracking
                if paragraph.style.name.startswith('Heading'):
                    current_section = paragraph.text
                
                # Analyze paragraph style
                paragraph_issues = self.analyze_paragraph(
                    paragraph, self.current_page, current_section
                )
                issues.extend(paragraph_issues)
                
            elif isinstance(element, Table):
                table = Table(element, document)
                
                # Analyze table style
                table_issues = self.analyze_table(
                    table, self.current_page, current_section
                )
                issues.extend(table_issues)
            
            progress.update()

        progress.complete()
        return issues