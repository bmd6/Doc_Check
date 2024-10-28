# doc_analyzer.py
import sys
from pathlib import Path
from typing import List, Tuple, Dict, Optional
from enum import Enum

from docx import Document

from analyzers.content_analyzer import ContentAnalyzer
from analyzers.style_analyzer import StyleAnalyzer
from analyzers.acronym_analyzer import AcronymAnalyzer
from models.issues import Issue, StyleIssue
from reports.text_reporter import TextReporter
from reports.csv_reporter import CSVReporter
from reports.org_reporter import OrgReporter
from reports.markdown_reporter import MarkdownReporter
from utils.logger import setup_logger
from utils.config_loader import ConfigLoader
from utils.progress import ProgressTracker

logger = setup_logger(__name__)

class OutputFormat(Enum):
    """Supported output formats for analysis reports."""
    TXT = "txt"
    CSV = "csv"
    ORG = "org"
    MD = "md"

class DocumentAnalyzer:
    """
    Main document analyzer class that coordinates all analysis components.
    
    This class:
    1. Loads and validates configuration
    2. Initializes all analyzers
    3. Coordinates the document analysis process
    4. Generates reports in requested formats
    """
    
    def __init__(self, config_path: str, acronyms_csv: Optional[str] = None, excluded_acronyms_csv: Optional[str] = None):
        """
        Initialize the document analyzer.
        
        Args:
            config_path: Path to configuration file
            acronyms_csv: Optional path to CSV file with known acronyms
        """
        logger.info("Initializing document analyzer...")
        
        # Load configuration
        self.config = ConfigLoader.load_config(config_path)
        
        # Initialize analyzers
        self.content_analyzer = ContentAnalyzer(self.config)
        self.style_analyzer = StyleAnalyzer()
        self.acronym_analyzer = None  # Will be initialized per document
        self.acronyms_csv = acronyms_csv
        self.excluded_acronyms_csv = excluded_acronyms_csv
        
        logger.info("Document analyzer initialized successfully")

    def analyze_document(self, doc_path: str) -> Tuple[List[Issue], List[StyleIssue]]:
        """
        Analyze a document for issues.
        
        Args:
            doc_path: Path to the document to analyze
            
        Returns:
            Tuple of (content issues, style issues)
            
        Raises:
            FileNotFoundError: If document doesn't exist
            ValueError: If document path is invalid
        """
        if not doc_path:
            raise ValueError("Document path is required.")

        logger.info(f"Starting analysis of document: {doc_path}")
        
        try:
            # Load document
            doc = Document(doc_path)
            logger.info(f"Document loaded successfully. Found {len(doc.paragraphs)} paragraphs.")
            
            # Initialize acronym analyzer for this document
            self.acronym_analyzer = AcronymAnalyzer(doc_path, self.acronyms_csv, self.excluded_acronyms_csv)
            
            # Track progress
            total_elements = len(doc.paragraphs) + sum(1 for _ in doc.tables)
            progress = ProgressTracker(total_elements, "Analyzing document")
            
            # Analyze document
            content_issues = []
            style_issues = []
            current_section = None
            current_page = 1
            chars_on_page = 0
            chars_per_page = 1800  # Approximate
            
            # Process each paragraph
            for paragraph in doc.paragraphs:
                # Update page tracking
                chars_on_page += len(paragraph.text)
                if chars_on_page > chars_per_page:
                    current_page += 1
                    chars_on_page = len(paragraph.text)
                
                # Update section tracking
                if paragraph.style.name.startswith('Heading'):
                    current_section = paragraph.text
                
                # Analyze content
                content_issues.extend(
                    self.content_analyzer.analyze_paragraph(
                        paragraph, current_page, current_section
                    )
                )
                
                # Analyze style
                style_issues.extend(
                    self.style_analyzer.analyze_paragraph(
                        paragraph, current_page, current_section
                    )
                )
                
                progress.update()
            
            # Process tables
            for table in doc.tables:
                # Update page count for tables
                table_text = "".join(
                    cell.text for row in table.rows for cell in row.cells
                )
                chars_on_page += len(table_text)
                if chars_on_page > chars_per_page:
                    current_page += 1
                    chars_on_page = len(table_text)
                
                # Analyze table style
                style_issues.extend(
                    self.style_analyzer.analyze_table(
                        table, current_page, current_section
                    )
                )
                
                progress.update()
            
            progress.complete()
            
            # Analyze acronyms
            logger.info("Analyzing acronyms...")
            self.acronym_analyzer.analyze_document()
            
            return content_issues, style_issues
            
        except Exception as e:
            logger.error(f"Error analyzing document: {e}")
            raise

    def generate_reports(self, content_issues: List[Issue], 
                        style_issues: List[StyleIssue], output_formats: List[str]) -> None:
        """
        Generate analysis reports in specified formats.
        
        Args:
            content_issues: List of content-related issues
            style_issues: List of style-related issues
            output_formats: List of desired output formats
        """
        logger.info("Generating analysis reports...")
        
        # Map formats to reporter classes
        reporters = {
            OutputFormat.TXT: TextReporter,
            OutputFormat.CSV: CSVReporter,
            OutputFormat.ORG: OrgReporter,
            OutputFormat.MD: MarkdownReporter
        }
        
        for format_str in output_formats:
            try:
                format_enum = OutputFormat(format_str.lower())
                reporter_class = reporters[format_enum]
                reporter = reporter_class(
                    self.content_analyzer.summary,
                    self.style_analyzer.font_usage,
                    self.acronym_analyzer
                )
                
                output_path = f"analysis_results.{format_str}"
                reporter.generate_report(content_issues, style_issues, output_path)
                logger.info(f"Generated {format_str} report: {output_path}")
                
            except (ValueError, KeyError) as e:
                logger.error(f"Unsupported output format: {format_str}")
                continue

    def print_summary(self, content_issues: List[Issue], 
                     style_issues: List[StyleIssue]) -> None:
        """
        Print analysis summary to console.
        
        Args:
            content_issues: List of content-related issues
            style_issues: List of style-related issues
        """
        print("\nAnalysis Summary")
        print("-" * 30)

        # Print Font Usage Summary
        print("\nFont Usage Analysis")
        print(self.style_analyzer.font_usage.get_formatted_summary())

        # Print Content Issues Summary
        print("\nContent Issues:")
        print(f"- Pronouns found: {self.content_analyzer.summary.pronouns}")
        print(f"- Contractions found: {self.content_analyzer.summary.contractions}")
        print("- Terminology conflicts:")
        for term, count in self.content_analyzer.summary.terminology.items():
            print(f"  - {term}: {count} occurrences")

        # Print Style Issues Summary
        print("\nStyle Issues:")
        print(f"- Font inconsistencies: {len([i for i in style_issues if i.type == 'Inconsistent Font'])}")
        print(f"- Caption issues: {len([i for i in style_issues if 'Caption' in i.type])}")
        print(f"- Header issues: {len([i for i in style_issues if 'Header' in i.type])}")
        print(f"- Table issues: {len([i for i in style_issues if 'Table' in i.type])}")

        # Print Acronym Summary
        if self.acronym_analyzer:
            stats = self.acronym_analyzer.get_statistics()
            print("\nAcronyms Summary:")
            print(f"- Total acronyms found: {stats['total_acronyms']}")
            print(f"- Defined acronyms: {stats['defined_acronyms']}")
            print(f"- Undefined acronyms: {stats['undefined_acronyms']}")

def main():
    """Main entry point for the document analyzer."""
    if len(sys.argv) < 2:
        logger.error("No document path provided.")
        print("Usage: python doc_analyzer.py <document_path> [config_path] [acronyms_csv]")
        sys.exit(1)

    try:
        # Get command line arguments
        doc_path = sys.argv[1]
        config_path = sys.argv[2] if len(sys.argv) > 2 else "config.json"
        acronyms_csv = sys.argv[3] if len(sys.argv) > 3 else None
        excluded_acronyms_csv = sys.argv[4] if len(sys.argv) > 4 else None

        # Print header
        print("\nDocument Analysis Tool")
        print("=" * 50)

        # Initialize analyzer
        analyzer = DocumentAnalyzer(config_path, acronyms_csv, excluded_acronyms_csv)

        # Analyze document
        content_issues, style_issues = analyzer.analyze_document(doc_path)

        # Generate reports in configured formats
        analyzer.generate_reports(
            content_issues, 
            style_issues, 
            analyzer.config.get("output_format", ["txt"])
        )

        # Print summary to console
        analyzer.print_summary(content_issues, style_issues)

        print("\nAnalysis complete.")
        print("=" * 50)
        
    except Exception as e:
        logger.error(f"Error during analysis: {e}")
        print(f"\nError: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()